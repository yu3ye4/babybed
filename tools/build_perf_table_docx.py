from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"C:\Users\18938\Desktop\报告\1.4主要性能指标_规范版.docx"


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tbl_pr.append(borders)


ROWS = [
    ["温湿度信号采集周期", "2 s", "软件可配置，建议 1~60 s", "1 s", "当前系统约每 2 s 刷新一次环境数据，兼顾实时性与通信开销"],
    ["压电信号采样频率", "50 Hz", "1~128 Hz", "1 Hz", "用于 PVDF 呼吸信号采样，当前默认配置为 50 Hz"],
    ["摄像头刷新率", "典型 30 fps", "由摄像头硬件和驱动配置决定", "按驱动档位调整", "建议最终以摄像头实测帧率为准"],
    ["数据上传周期", "2 s", "软件可配置，建议 1~60 s", "1 s", "影响页面刷新频率与网络带宽占用"],
    ["温度测量范围", "-40~85 ℃", "不可调", "精度 ±0.3 ℃（典型）", "用于环境温度采集与有效性检查"],
    ["湿度测量范围", "0~100 %RH", "不可调", "精度 ±2 %RH（典型）", "用于环境湿度采集与有效性检查"],
    ["环境温度阈值", "18.00~30.00 ℃", "软件可配置", "0.01 ℃", "通过 temp_min、temp_max 远程配置环境报警阈值"],
    ["环境湿度阈值", "30.00~70.00 %RH", "软件可配置", "0.01 %RH", "通过 humi_min、humi_max 远程配置环境报警阈值"],
    ["温度严重告警阈值", "<15.00 ℃ 或 >35.00 ℃", "软件中固定，可按需求修改", "0.01 ℃", "超过该范围直接进入高风险环境告警"],
    ["湿度严重告警阈值", ">80.00 %RH", "软件中固定，可按需求修改", "0.01 %RH", "用于高湿环境快速预警"],
    ["潮湿怀疑判据", "湿度增量 ≥ 8.00 %RH 且持续 30 s", "软件可配置", "0.01 %RH；1 s", "用于识别婴儿床受潮或尿湿风险"],
    ["PVDF 环形缓存长度", "500 点", "1~500 点", "1 点", "50 Hz 下可覆盖约 10 s 波形数据"],
    ["呼吸统计默认窗口", "300 点（约 6 s）", "1~500 点", "1 点", "用于峰峰值、能量和周期性统计"],
    ["呼吸活动窗口", "100 点（约 2 s）", "1~500 点", "1 点", "用于短时呼吸活动辅助判断"],
    ["压电峰峰值阈值", "10 mV；150 mV", "软件可配置", "1 mV", "分别用于基础呼吸变化与明显活动识别"],
    ["压电能量阈值", "55 mV", "软件可配置", "1 mV", "作为活跃度判断辅助量"],
    ["呼吸周期性判据", "周期峰谷次数 2~12 次", "软件可配置", "1 次", "用于判断波形是否具有稳定呼吸节律"],
    ["弱呼吸判定时长", "连续 5 s", "软件可配置", "1 s", "连续弱活动时判定为弱呼吸风险"],
    ["疑似呼吸暂停判定时长", "连续 10 s", "软件可配置", "1 s", "连续较长时间无有效周期信号时触发高风险提示"],
    ["啼哭开始判定阈值", "crying ≥ 0.80", "0~1", "0.01", "哭声类别置信度达到 0.80 时判定为啼哭开始"],
    ["啼哭停止判定阈值", "crying < 0.55 连续 4 窗口", "阈值 0~1；窗口数为正整数", "0.01；1 个窗口", "用于抑制短时波动，降低误停概率"],
    ["哭声模型输入参数", "16 kHz，16000 点窗口，2 类输出", "由模型结构决定", "采样间隔 0.0625 ms", "当前模型输出 crying/silence 两类结果"],
    ["睡姿识别输出", "4 类", "由模型类别定义决定", "置信度步长 0.01", "当前仓库可确认类别为 safe_sleep、side_sleep、near_edge、blanket_cover"],
    ["视觉识别置信度阈值", "默认 0.25", "0~1", "0.01", "低于阈值的检测结果不参与输出"],
    ["历史缓存记录条数", "500 条", "1~500 条", "1 条", "用于 Web Console 历史记录与呼吸波形缓存"],
]


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run("1.4 主要性能指标（规范版）")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p2 = doc.add_paragraph(
        "表头建议采用“性能指标 | 当前设定值 | 可调范围/量程 | 调节步长/分辨率/精度 | 说明”，以增强参数表述的严谨性。"
    )
    for run in p2.runs:
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    headers = ["性能指标", "当前设定值", "可调范围/量程", "调节步长/分辨率/精度", "说明"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)

    for i, text in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], text, bold=True, size=10.5)

    for row in ROWS:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, size=9.5)
            if i in (0, 4):
                cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    widths = [Cm(3.5), Cm(2.6), Cm(4.0), Cm(3.8), Cm(5.2)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    p3 = doc.add_paragraph("表1.1 监测系统主要性能指标及参数配置范围")
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p3.runs:
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
